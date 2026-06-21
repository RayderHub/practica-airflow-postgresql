from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "entrega_practica_ia_datos" / "analisis_comparativo_casos.docx"


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = None
    return paragraph


def add_paragraph(document, text=""):
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(3)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for paragraph in cells[index].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
    document.add_paragraph()
    return table


def add_cover(document):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Analisis comparativo de IA, Machine Learning, Data Mining y Big Data")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Casos de estudio y aplicacion a la practica ETL")
    run.font.name = "Arial"
    run.font.size = Pt(12)

    document.add_paragraph()
    add_paragraph(document, "Actividad: Analisis comparativo y planeacion de analisis de datos")
    add_paragraph(document, "Alumno: Rayder")
    add_paragraph(document, "Repositorio de apoyo: https://github.com/RayderHub/practica-airflow-postgresql")
    document.add_page_break()


def build_document():
    document = Document()
    set_normal_style(document)
    add_cover(document)

    add_heading(document, "1. Analisis comparativo", 1)
    add_paragraph(
        document,
        "En esta seccion se comparan cuatro conceptos que se usan mucho en proyectos de datos. "
        "Aunque estan relacionados, no significan lo mismo. La inteligencia artificial es el campo mas amplio; "
        "machine learning es una forma de aprender a partir de datos; data mining ayuda a encontrar patrones; "
        "y big data se enfoca en manejar informacion de gran volumen o complejidad.",
    )

    add_table(
        document,
        ["Concepto", "Caracteristicas", "Beneficios", "Restricciones y retos", "Casos de aplicacion", "Lenguajes y herramientas"],
        [
            [
                "Inteligencia artificial",
                "Busca que un sistema realice tareas asociadas con inteligencia humana, como razonar, decidir, interpretar texto o reconocer patrones.",
                "Automatiza procesos, mejora decisiones y permite crear servicios personalizados.",
                "Necesita datos confiables, puede tener sesgos y a veces sus resultados son dificiles de explicar.",
                "Chatbots, recomendaciones, deteccion de fraudes, vision por computadora y asistentes inteligentes.",
                "Python, Java, TensorFlow, PyTorch, Keras, Azure AI y Google Vertex AI.",
            ],
            [
                "Machine learning",
                "Es una rama de la IA que usa datos para entrenar modelos capaces de predecir, clasificar o agrupar informacion.",
                "Permite anticipar comportamientos, clasificar clientes y mejorar con nuevos datos.",
                "Requiere datos limpios, variables utiles y una evaluacion correcta para evitar sobreajuste.",
                "Prediccion de abandono de clientes, ventas, riesgo crediticio, spam y mantenimiento predictivo.",
                "Python, R, Scikit-learn, XGBoost, Pandas, NumPy, Jupyter y MLflow.",
            ],
            [
                "Data mining",
                "Se enfoca en descubrir patrones, relaciones o tendencias dentro de datos historicos.",
                "Convierte datos almacenados en informacion util para tomar decisiones.",
                "Los patrones encontrados deben interpretarse con cuidado, porque correlacion no siempre significa causalidad.",
                "Segmentacion de clientes, canasta de mercado, deteccion de anomalias y analisis de comportamiento.",
                "SQL, Python, R, Weka, RapidMiner, Orange, PostgreSQL, Power BI y Tableau.",
            ],
            [
                "Big data",
                "Trabaja con datos de gran volumen, velocidad o variedad, muchas veces mediante procesamiento distribuido.",
                "Permite analizar informacion masiva de multiples fuentes y casi en tiempo real.",
                "Puede ser costoso, requiere infraestructura y necesita buen gobierno de datos.",
                "Analisis de logs, sensores IoT, redes sociales, transacciones masivas y sistemas de recomendacion.",
                "Python, Scala, Java, Spark, Hadoop, Kafka, Hive, Databricks, BigQuery y Snowflake.",
            ],
        ],
    )

    add_paragraph(
        document,
        "En un proyecto real, estos conceptos pueden trabajar juntos. Big data ayuda a almacenar y procesar grandes cantidades de datos; "
        "data mining permite entender patrones; machine learning crea modelos predictivos; y la inteligencia artificial usa esos modelos "
        "para apoyar o automatizar decisiones.",
    )

    add_heading(document, "2. Caso de estudio 1: Electro-Hogar S.A.", 1)
    add_heading(document, "2.1 Objetivo y alcance", 2)
    add_paragraph(
        document,
        "Electro-Hogar S.A. es una cadena minorista de electronica y electrodomesticos que ha detectado una baja en la retencion de clientes. "
        "El problema principal es que muchos clientes compran una vez, pero no regresan o dejan de comprar despues de poco tiempo.",
    )
    add_paragraph(document, "Objetivo general:")
    add_paragraph(
        document,
        "Analizar el comportamiento historico de los clientes para identificar factores asociados al abandono, mejorar la retencion y aumentar el valor de vida del cliente.",
    )
    add_paragraph(document, "Objetivos especificos:")
    add_bullets(
        document,
        [
            "Identificar patrones de clientes que no realizan una segunda compra.",
            "Segmentar clientes por recencia, frecuencia y monto de compra.",
            "Detectar variables relacionadas con abandono o baja lealtad.",
            "Proponer acciones de retencion basadas en datos.",
            "Preparar la informacion para una fase posterior de machine learning.",
        ],
    )
    add_paragraph(
        document,
        "El alcance inicial incluye datos de ventas, clientes, productos, promociones, canales de venta, ubicacion y servicio postventa. "
        "No se considera todavia la puesta en produccion de un modelo predictivo, porque primero se necesita diagnosticar la calidad de datos y definir indicadores.",
    )

    add_heading(document, "2.2 Justificacion de la metodologia", 2)
    add_paragraph(
        document,
        "La metodologia mas adecuada para este caso es CRISP-DM, porque organiza el trabajo en etapas claras: entendimiento del negocio, "
        "entendimiento de los datos, preparacion, modelado, evaluacion y despliegue. Esta metodologia ayuda a no empezar directamente con modelos, "
        "sino primero entender que significa abandono para la empresa y que datos realmente sirven para analizarlo.",
    )
    add_paragraph(
        document,
        "Tambien conviene usar analisis RFM, ya que permite clasificar clientes con base en recencia, frecuencia y valor monetario. "
        "Esto facilita detectar clientes activos, clientes en riesgo, clientes de alto valor y clientes que probablemente ya se perdieron.",
    )

    add_heading(document, "2.3 Planeacion de etapas", 2)
    add_table(
        document,
        ["Etapa", "Actividades principales"],
        [
            ["Entendimiento del negocio", "Definir que se considera abandono, establecer periodo de analisis y seleccionar indicadores como recompra, ticket promedio y valor de vida del cliente."],
            ["Entendimiento de datos", "Revisar fuentes disponibles, detectar valores faltantes, duplicados y formatos inconsistentes."],
            ["Preparacion de datos", "Integrar ventas y clientes, homologar catalogos y crear variables como dias desde ultima compra, total gastado y numero de compras."],
            ["Analisis exploratorio", "Comparar clientes retenidos y no retenidos por categoria, canal, region y periodo."],
            ["Segmentacion", "Aplicar RFM y clustering para encontrar grupos de comportamiento."],
            ["Modelado posterior", "Preparar una variable objetivo de abandono y separar datos en entrenamiento y prueba."],
            ["Recomendaciones", "Proponer campanas personalizadas, promociones de segunda compra y estrategias de fidelizacion."],
        ],
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "3. Caso de estudio 2: Aplicacion a la practica ETL", 1)
    add_heading(document, "3.1 Fase documental y diagnostico inicial", 2)
    add_paragraph(
        document,
        "La practica ETL trabajada previamente usa un dataset transaccional de e-commerce. Aunque la instruccion menciona contrataciones reales, "
        "el archivo disponible para esta practica corresponde a ventas. Por eso, el caso se adapta al contexto de ventas, respetando la misma logica de diagnostico, "
        "limpieza, diseno dimensional y preparacion para mineria de datos.",
    )
    add_paragraph(document, "Diagnostico inicial del archivo crudo:")
    add_bullets(
        document,
        [
            "Filas crudas: 541,909.",
            "Columnas: 8.",
            "Clientes faltantes en CustomerID: 135,080.",
            "Descripciones faltantes: 1,454.",
            "Facturas canceladas detectadas por prefijo C: 9,288.",
            "Paises distintos: 38.",
            "Registros duplicados exactos: 5,268.",
            "Cantidades negativas: 10,624.",
            "Precios unitarios menores o iguales a cero: 2,517.",
        ],
    )
    add_paragraph(
        document,
        "Estos datos muestran que el archivo no debe usarse directamente para mineria de datos. Primero se necesita limpiar textos, validar formatos, "
        "tratar valores faltantes y separar las transacciones validas de las cancelaciones o registros inconsistentes.",
    )

    add_heading(document, "3.2 Grano de la tabla de hechos", 2)
    add_paragraph(
        document,
        "El grano seleccionado para la tabla de hechos es una linea de producto dentro de una factura de venta. Es decir, cada registro representa "
        "un producto especifico vendido en una factura, asociado a una fecha, cliente, pais, cantidad y precio.",
    )
    add_paragraph(
        document,
        "Este grano es adecuado porque conserva el mayor detalle disponible. Permite analizar ventas por producto, cliente, pais y fecha, "
        "ademas de calcular indicadores como importe por linea, cantidad vendida, devoluciones y ventas por periodo. Si se usara un nivel mas agregado, "
        "como factura o cliente, se perderia informacion importante para analisis de canasta de mercado o modelos predictivos.",
    )

    add_heading(document, "3.3 Esquema de data warehouse", 2)
    add_table(
        document,
        ["Tabla", "Tipo", "Contenido principal"],
        [
            ["fact_venta_linea", "Hechos", "id_fecha, id_cliente, id_producto, id_pais, invoice_no, quantity, unit_price, importe_linea, es_cancelacion, alto_valor_transaccion."],
            ["dim_fecha", "Dimension", "fecha, dia, mes, anio, trimestre y dia_semana."],
            ["dim_cliente", "Dimension", "customer_id_original, segmento_cliente y bandera de cliente identificado."],
            ["dim_producto", "Dimension", "stock_code, descripcion_limpia y categoria estimada."],
            ["dim_pais", "Dimension", "pais_limpio y region."],
            ["dim_factura", "Dimension", "invoice_no, tipo_factura y bandera de cancelacion."],
        ],
    )

    add_heading(document, "3.4 Tecnicas de limpieza de datos", 2)
    add_bullets(
        document,
        [
            "Estandarizacion de texto en mayusculas.",
            "Eliminacion de caracteres especiales o corruptos.",
            "Normalizacion de espacios multiples.",
            "Reemplazo de descripciones nulas por SIN_DESCRIPTION.",
            "Validacion de facturas con expresiones regulares.",
            "Conversion correcta de fechas, cantidades y precios.",
            "Calculo de importe_linea como Quantity por UnitPrice.",
            "Identificacion de cancelaciones mediante el prefijo C en InvoiceNo.",
            "Filtrado de ventas positivas para crear archivos de mineria de datos.",
        ],
    )

    add_heading(document, "3.5 Expresiones regulares propuestas", 2)
    add_table(
        document,
        ["Objetivo", "Expresion regular", "Uso"],
        [
            ["Limpiar descripciones", r"[^A-Z0-9\s\-]", "Elimina simbolos y caracteres corruptos, conservando letras, numeros, espacios y guiones."],
            ["Normalizar espacios", r"\s+", "Sustituye varios espacios por uno solo."],
            ["Validar facturas", r"^C?\d{6}$", "Acepta facturas de seis digitos y cancelaciones que inician con C."],
            ["Eliminar variantes corporativas", r"\b(SA DE CV|S\.A\. DE C\.V\.|SAPI DE CV|S\. DE R\.L\.|SRL|SC|AC)\b", "Sirve para homologar nombres de proveedores si el dataset fuera de contrataciones."],
            ["Limpiar dependencias o proveedores", r"[^A-Z0-9\s&]", "Quita caracteres no necesarios en nombres institucionales."],
        ],
    )

    add_heading(document, "3.6 Variable objetivo", 2)
    add_paragraph(
        document,
        "Se creo la variable binaria alto_valor_transaccion. Esta variable vale 1 cuando el importe de la linea de venta es igual o mayor a la mediana "
        "de importes positivos, y vale 0 cuando el importe es menor. Esta definicion permite preparar un problema de clasificacion para una etapa posterior "
        "de mineria de datos o machine learning.",
    )

    add_heading(document, "3.7 Particionamiento Train/Test", 2)
    add_paragraph(
        document,
        "Para demostrar que los datos estan listos para la siguiente fase, se genero un subconjunto balanceado y se exportaron dos archivos CSV: "
        "train_ecommerce_balanced.csv y test_ecommerce_balanced.csv.",
    )
    add_table(
        document,
        ["Archivo", "Total de registros", "Clase 0", "Clase 1"],
        [
            ["train_ecommerce_balanced.csv", "16,000", "8,000", "8,000"],
            ["test_ecommerce_balanced.csv", "4,000", "2,000", "2,000"],
        ],
    )
    add_paragraph(
        document,
        "El resultado muestra una distribucion balanceada en ambos archivos. Esto ayuda a evitar que un modelo posterior aprenda mas de una clase que de otra.",
    )

    add_heading(document, "4. Archivos generados para la entrega", 1)
    add_bullets(
        document,
        [
            "analisis_comparativo_casos.docx: documento principal de la actividad.",
            "train_ecommerce_balanced.csv: datos de entrenamiento balanceados.",
            "test_ecommerce_balanced.csv: datos de prueba balanceados.",
            "github_link.txt: enlace al repositorio de apoyo.",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
